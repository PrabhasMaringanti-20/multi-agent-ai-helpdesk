"""DOCX text extraction (python-docx, imported lazily)."""

from __future__ import annotations

import io

from app.core.exceptions import ValidationError


def parse_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ValidationError("python-docx is not installed.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ValidationError(f"Failed to parse DOCX: {exc}") from exc

    parts: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    if not text.strip():
        raise ValidationError("DOCX contained no extractable text.")
    return text


__all__ = ["parse_docx"]
