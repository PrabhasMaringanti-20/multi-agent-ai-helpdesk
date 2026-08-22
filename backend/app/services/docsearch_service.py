"""Document-Intelligence service — parse, index and AI-search uploaded files.

Handles PDF / Word / text / Excel (one chosen tab) / URL. Each source is split
into passages tagged with a human-readable LOCATION (page / sheet+row / section)
and the VERBATIM text. Search uses Postgres full-text; the LLM then writes a
short readable "thumbnail" for the top hits (one call per search — quota-safe),
while the card can be opened to reveal the exact original passage.
"""

from __future__ import annotations

import io
import ipaddress
import re
import socket
import uuid
from typing import Any
from urllib.parse import urlparse

import httpx
from pydantic import BaseModel
from sqlalchemy import func, select

from app.models.docsearch import UploadedChunk, UploadedDocument
from app.providers.base import ChatMessage
from app.repositories.kb_repo import _or_tsquery_terms

_MAX_CHUNKS = 4000
_CHUNK = 1000


def _split(text: str, size: int = _CHUNK) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]
    out: list[str] = []
    cur = ""
    for para in re.split(r"\n\s*\n", text):
        if cur and len(cur) + len(para) > size:
            out.append(cur.strip())
            cur = ""
        cur += para + "\n\n"
        while len(cur) > size * 2:
            out.append(cur[:size].strip())
            cur = cur[size:]
    if cur.strip():
        out.append(cur.strip())
    return [p for p in out if p]


def sheet_names(data: bytes) -> list[str]:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    return list(wb.sheetnames)


def _parse_pdf(data: bytes) -> list[tuple[str, str]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    passages: list[tuple[str, str]] = []
    for i, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        for j, chunk in enumerate(_split(text)):
            loc = f"Page {i}" + (f" · part {j + 1}" if j else "")
            passages.append((loc, chunk))
    return passages


def _parse_docx(data: bytes) -> list[tuple[str, str]]:
    from docx import Document

    doc = Document(io.BytesIO(data))
    full = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [(f"Section {j}", c) for j, c in enumerate(_split(full), 1)]


def _parse_excel(data: bytes, sheet: str | None) -> tuple[str, list[tuple[str, str]]]:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    name = sheet if sheet in wb.sheetnames else wb.sheetnames[0]
    ws = wb[name]
    header: list[str] = []
    passages: list[tuple[str, str]] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        cells = ["" if c is None else str(c).strip() for c in row]
        if i == 0:
            header = [c.replace("\n", " ").strip() for c in cells]
            continue
        if not any(cells):
            continue
        pairs = [
            f"{(header[k] if k < len(header) and header[k] else f'col{k + 1}')}: {v}"
            for k, v in enumerate(cells)
            if v
        ]
        text = "\n".join(pairs)  # one field per line -> readable in the drill-down
        if text:
            passages.append((f"Sheet {name} · Row {i + 1}", text))
    return name, passages


def _parse_text(data: bytes) -> list[tuple[str, str]]:
    text = data.decode("utf-8", errors="ignore")
    return [(f"Part {j}", c) for j, c in enumerate(_split(text), 1)]


def _url_is_safe(url: str) -> bool:
    try:
        p = urlparse(url)
        if p.scheme not in ("http", "https") or not p.hostname:
            return False
        host = p.hostname.lower()
        if host in ("localhost", "127.0.0.1", "::1"):
            return False
        try:
            ip = ipaddress.ip_address(socket.gethostbyname(host))
            if ip.is_private or ip.is_loopback or ip.is_link_local:
                return False
        except (socket.gaierror, ValueError):
            pass
        return True
    except Exception:  # noqa: BLE001
        return False


def fetch_url(url: str) -> list[tuple[str, str]]:
    if not _url_is_safe(url):
        raise ValueError("Only public http/https URLs are allowed.")
    resp = httpx.get(
        url, timeout=15, follow_redirects=True, headers={"User-Agent": "HelpdeskDocSearch/1.0"}
    )
    resp.raise_for_status()
    html = resp.text
    html = re.sub(r"<(script|style|noscript)[^>]*>.*?</\1>", " ", html, flags=re.S | re.I)
    text = re.sub(r"<[^>]+>", " ", html)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n\s*", "\n\n", text)
    return [(f"Section {j}", c) for j, c in enumerate(_split(text), 1)]


class _Summaries(BaseModel):
    summaries: list[str] = []


class DocSearchService:
    def __init__(self, session: Any, llm: Any) -> None:
        self.session = session
        self.llm = llm

    # ---- ingest ----------------------------------------------------------- #
    async def ingest(
        self,
        principal: Any,
        *,
        filename: str,
        source_type: str,
        passages: list[tuple[str, str]],
        source_ref: str | None = None,
        sheet: str | None = None,
    ) -> dict[str, Any]:
        passages = passages[:_MAX_CHUNKS]
        doc = UploadedDocument(
            id=uuid.uuid4(),
            org_id=principal.org_id,
            user_id=principal.user_id,
            filename=filename,
            source_type=source_type,
            source_ref=source_ref,
            sheet=sheet,
            chunk_count=len(passages),
        )
        self.session.add(doc)
        for idx, (location, text) in enumerate(passages):
            self.session.add(
                UploadedChunk(
                    id=uuid.uuid4(),
                    doc_id=doc.id,
                    org_id=principal.org_id,
                    user_id=principal.user_id,
                    chunk_index=idx,
                    location=location,
                    text=text,
                )
            )
        await self.session.commit()
        return {
            "id": str(doc.id),
            "filename": filename,
            "source_type": source_type,
            "sheet": sheet,
            "chunk_count": len(passages),
        }

    async def list_documents(self, principal: Any) -> list[dict[str, Any]]:
        stmt = (
            select(UploadedDocument)
            .where(
                UploadedDocument.org_id == principal.org_id,
                UploadedDocument.user_id == principal.user_id,
            )
            .order_by(UploadedDocument.created_at.desc())
        )
        rows = (await self.session.execute(stmt)).scalars().all()
        return [
            {
                "id": str(d.id),
                "filename": d.filename,
                "source_type": d.source_type,
                "sheet": d.sheet,
                "chunk_count": d.chunk_count,
            }
            for d in rows
        ]

    async def delete_document(self, principal: Any, doc_id: uuid.UUID) -> bool:
        doc = await self.session.get(UploadedDocument, doc_id)
        if doc is None or doc.org_id != principal.org_id or doc.user_id != principal.user_id:
            return False
        await self.session.delete(doc)
        await self.session.commit()
        return True

    # ---- search ----------------------------------------------------------- #
    async def search(self, principal: Any, query: str, limit: int = 8) -> dict[str, Any]:
        terms = _or_tsquery_terms(query)
        if not terms:
            return {"query": query, "hits": []}
        tsquery = func.to_tsquery("english", terms)
        rank = func.ts_rank(UploadedChunk.text_fts, tsquery).label("rank")
        stmt = (
            select(UploadedChunk, UploadedDocument.filename, UploadedDocument.source_type, rank)
            .join(UploadedDocument, UploadedChunk.doc_id == UploadedDocument.id)
            .where(
                UploadedChunk.org_id == principal.org_id,
                UploadedChunk.user_id == principal.user_id,
                UploadedChunk.text_fts.op("@@")(tsquery),
            )
            .order_by(rank.desc())
            .limit(limit)
        )
        rows = (await self.session.execute(stmt)).all()
        hits = [
            {
                "chunk_id": str(c.id),
                "document_id": str(c.doc_id),
                "filename": fname,
                "source_type": stype,
                "location": c.location,
                "text": c.text,  # verbatim, for drill-down
                "score": round(float(r or 0.0), 4),
                "summary": "",
            }
            for (c, fname, stype, r) in rows
        ]

        if hits:
            summaries = await self._summarize(query, hits[:5])
            for i, s in enumerate(summaries):
                if i < len(hits):
                    hits[i]["summary"] = s
        for h in hits:
            if not h["summary"]:
                clean = re.sub(r"\s+", " ", h["text"]).replace(" | ", " · ").strip()
                h["summary"] = (clean[:200] + "…") if len(clean) > 200 else clean
        return {"query": query, "hits": hits}

    async def _summarize(self, query: str, hits: list[dict[str, Any]]) -> list[str]:
        numbered = "\n\n".join(
            f"[{i + 1}] ({h['filename']} · {h['location']}): {h['text'][:500]}"
            for i, h in enumerate(hits)
        )
        messages = [
            ChatMessage(
                role="system",
                content=(
                    "For each numbered passage, write ONE short plain-language sentence describing "
                    "what it says that is relevant to the user's query. Do not add facts. Return "
                    'JSON {"summaries": ["...", "..."]} with one entry per passage, in order.'
                ),
            ),
            ChatMessage(role="user", content=f"Query: {query}\n\nPassages:\n{numbered}"),
        ]
        try:
            res: _Summaries = await self.llm.generate_structured(messages, _Summaries)
            return list(res.summaries)
        except Exception:  # noqa: BLE001 - fall back to raw snippets
            return []


class UploadsSearcher:
    """Org-wide full-text search over uploaded files, returning ``RetrievedChunk``
    objects so the AI chat can answer from documents anyone in the org attached.
    """

    def __init__(self, session: Any) -> None:
        self.session = session

    async def search(self, org_id: Any, query: str, limit: int = 4) -> list[Any]:
        from app.agents.state import RetrievedChunk

        terms = _or_tsquery_terms(query or "")
        if not terms:
            return []
        tsquery = func.to_tsquery("english", terms)
        rank = func.ts_rank(UploadedChunk.text_fts, tsquery)
        stmt = (
            select(UploadedChunk, UploadedDocument.filename, rank)
            .join(UploadedDocument, UploadedChunk.doc_id == UploadedDocument.id)
            .where(UploadedChunk.org_id == org_id, UploadedChunk.text_fts.op("@@")(tsquery))
            .order_by(rank.desc())
            .limit(limit)
        )
        try:
            rows = (await self.session.execute(stmt)).all()
        except Exception:  # noqa: BLE001 - never break chat if uploads search fails
            return []
        out = []
        for c, fname, _r in rows:
            out.append(
                RetrievedChunk(
                    chunk_id=str(c.id),
                    doc_id=str(c.doc_id),
                    text=c.text,
                    score=0.75,
                    rerank_score=0.75,  # an FTS hit in an uploaded file is a real match
                    source_uri=f"upload://{fname} · {c.location}",
                    category_key=None,
                )
            )
        return out


__all__ = [
    "DocSearchService",
    "UploadsSearcher",
    "sheet_names",
    "fetch_url",
    "_parse_pdf",
    "_parse_docx",
    "_parse_excel",
    "_parse_text",
]
